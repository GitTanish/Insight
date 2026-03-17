import os
import pandas as pd
import streamlit as st
import io
from config import ENCODINGS_TO_TRY, DELIMITERS_TO_TRY, MAX_ROWS, MAX_PLOT_FILES


def initialize_session_state():
    """Initialize session state variables."""
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'quick_query' not in st.session_state:
        st.session_state.quick_query = None
    if 'agent_initialized' not in st.session_state:
        st.session_state.agent_initialized = False
    if 'agent' not in st.session_state:
        st.session_state.agent = None
    if 'plot_counter' not in st.session_state:
        st.session_state.plot_counter = 0


def get_data_summary(df):
    """Get comprehensive data summary statistics."""
    summary = {
        "rows": len(df),
        "columns": len(df.columns),
        "memory_usage": df.memory_usage(deep=True).sum(),
        "numeric_columns": len(df.select_dtypes(include=['number']).columns),
        "categorical_columns": len(df.select_dtypes(include=['object']).columns),
        "missing_values": df.isnull().sum().sum(),
        "duplicate_rows": df.duplicated().sum()
    }
    return summary


def clear_conversation():
    """Clear conversation history and reset agent state."""
    st.session_state.history = []
    st.session_state.agent_initialized = False
    st.session_state.agent = None
    st.session_state.plot_counter = 0
    
    # Clean up any existing plot files
    for i in range(MAX_PLOT_FILES):
        plot_path = f"temp_plot_{i}.png"
        if os.path.exists(plot_path):
            os.remove(plot_path)
    
    st.rerun()


def validate_csv(uploaded_file):
    """Validate uploaded CSV file."""
    try:
        raw_data = uploaded_file.read()
    except Exception as e:
        return False, f"Failed to read uploaded file content: {str(e)}"
    
    df = None
    last_exception = None
    
    for encoding in ENCODINGS_TO_TRY:
        for delimiter in DELIMITERS_TO_TRY:
            try:
                df = pd.read_csv(io.BytesIO(raw_data), encoding=encoding, delimiter=delimiter)
                if not df.empty and len(df.columns) > 0:
                    if len(df) > MAX_ROWS:
                        return False, f"File is too large (max {MAX_ROWS:,} rows)"
                    return True, f"File validation successful (encoding: {encoding}, delimiter: '{delimiter}')"
                else:
                    last_exception = f"Loaded with encoding='{encoding}' and delimiter='{delimiter}', but DataFrame is empty or has no columns."
            except Exception as e:
                last_exception = f"Failed to load with encoding='{encoding}' and delimiter='{delimiter}': {e}"
    
    if df is None or df.empty or len(df.columns) == 0:
        return False, f"Could not parse CSV file. Last attempt failed with: {last_exception}"
    
    return False, "An unexpected error occurred during CSV validation."


def validate_api_key(api_key):
    """Validate Groq API key format."""
    if not api_key:
        return False, "API key is required"
    if not api_key.startswith('gsk_'):
        return False, "Invalid API key format (should start with 'gsk_')"
    return True, "API key format is valid"


def load_dataframe_from_session():
    """Load DataFrame from session state with error handling."""
    if 'uploaded_file_content' not in st.session_state or st.session_state.uploaded_file_content is None:
        return None, "No file uploaded"
    
    try:
        df = None
        last_load_error = None
        
        for encoding in ENCODINGS_TO_TRY:
            for delimiter in DELIMITERS_TO_TRY:
                try:
                    df = pd.read_csv(
                        io.BytesIO(st.session_state.uploaded_file_content), 
                        encoding=encoding, 
                        delimiter=delimiter
                    )
                    if not df.empty and len(df.columns) > 0:
                        return df, None
                    else:
                        last_load_error = f"Loaded with encoding='{encoding}' and delimiter='{delimiter}', but DataFrame is empty or has no columns."
                except Exception as e:
                    last_load_error = f"Failed to load with encoding='{encoding}' and delimiter='{delimiter}': {e}"
        
        if df is None or df.empty or len(df.columns) == 0:
            return None, f"Could not load DataFrame after trying multiple encodings and delimiters. Last error: {last_load_error}"
        
        return df, None
        
    except Exception as e:
        return None, f"Failed to process the CSV file: {str(e)}"


def cleanup_plot_files():
    """Clean up temporary plot files."""
    for i in range(MAX_PLOT_FILES):
        plot_path = f"temp_plot_{i}.png"
        if os.path.exists(plot_path):
            try:
                os.remove(plot_path)
            except Exception:
                pass  # Ignore cleanup errors


def generate_export_docx(history, dataset_name="Dataset", date_str=""):
    """Generate a DOCX file from the conversation history."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Please install python-docx to use export features.")
        
    doc = Document()
    
    title = f"Insight Analysis Report — {dataset_name}"
    if date_str:
        title += f" — {date_str}"
        
    doc.add_heading(title, 0)
    
    if not history:
        doc.add_paragraph("No conversation history available.")
    
    figure_count = 1
    
    for msg in history:
        role = "You" if msg["role"] == "user" else "Insight"
        doc.add_heading(role, level=2)
        doc.add_paragraph(msg["content"])
        
        if "plot_paths" in msg:
            for plot_path in msg["plot_paths"]:
                if os.path.exists(plot_path):
                    try:
                        doc.add_picture(plot_path, width=Inches(6.0))
                        
                        caption_para = doc.add_paragraph(f"Figure {figure_count}: Analysis Visualization")
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if caption_para.runs:
                            caption_run = caption_para.runs[0]
                            caption_run.italic = True
                            caption_run.font.size = Pt(9)
                            
                        figure_count += 1
                    except Exception:
                        doc.add_paragraph(f"[Image failed to load: {plot_path}]")
                        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()